from dataclasses import dataclass
from pathlib import Path
from typing import Optional
import csv
import zipfile
import xml.etree.ElementTree as ET
import subprocess

from fastapi import HTTPException


@dataclass
class ExtractedText:
    text: str
    source_type: str
    parse_method: str
    page_count: Optional[int] = None


def extract_text_from_local_file(file_path: str, content_type: str | None = None) -> ExtractedText:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return extract_txt(file_path)

    if suffix == ".csv":
        return extract_csv(file_path)

    if suffix == ".pdf":
        return extract_pdf(file_path)

    if suffix == ".docx":
        return extract_docx(file_path)

    if suffix in [".xlsx", ".xlsm"]:
        return extract_xlsx(file_path)

    if suffix == ".hwpx":
        return extract_hwpx(file_path)

    if suffix == ".hwp":
        return extract_hwp(file_path)

    if suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"]:
        return extract_image_with_ocr(file_path)

    raise HTTPException(400, f"지원하지 않는 파일 형식입니다: {suffix}")


def extract_txt(file_path: str) -> ExtractedText:
    encodings = ["utf-8", "utf-8-sig", "cp949", "euc-kr"]

    last_error = None

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()

            return ExtractedText(
                text=text or "",
                source_type="txt",
                parse_method=f"plain:{encoding}",
            )
        except UnicodeDecodeError as e:
            last_error = e
        except Exception as e:
            raise HTTPException(500, f"TXT 텍스트 추출 실패: {e}")

    raise HTTPException(500, f"TXT 인코딩 판별 실패: {last_error}")


def extract_csv(file_path: str) -> ExtractedText:
    encodings = ["utf-8", "utf-8-sig", "cp949", "euc-kr"]

    for encoding in encodings:
        try:
            lines = []

            with open(file_path, "r", encoding=encoding, newline="") as f:
                reader = csv.reader(f)

                for row_idx, row in enumerate(reader, start=1):
                    values = [str(cell).strip() for cell in row if str(cell).strip()]
                    if values:
                        lines.append(f"Row {row_idx}: " + " | ".join(values))

            return ExtractedText(
                text="\n".join(lines),
                source_type="csv",
                parse_method=f"csv:{encoding}",
            )

        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise HTTPException(500, f"CSV 텍스트 추출 실패: {e}")

    raise HTTPException(500, "CSV 인코딩 판별 실패")


def extract_pdf(file_path: str) -> ExtractedText:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise HTTPException(500, "PyMuPDF 라이브러리가 필요합니다. pip install pymupdf")

    try:
        doc = fitz.open(file_path)
        pages = []

        for page in doc:
            page_text = page.get_text("text") or ""
            pages.append(page_text.strip())

        merged = "\n\n".join([p for p in pages if p]).strip()
        page_count = len(doc)

        doc.close()

        if len(merged) < 30:
            return extract_pdf_with_ocr(file_path)

        return ExtractedText(
            text=merged,
            source_type="pdf",
            parse_method="pymupdf_text",
            page_count=page_count,
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"PDF 텍스트 추출 실패: {e}")


def extract_docx(file_path: str) -> ExtractedText:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(500, "python-docx 라이브러리가 필요합니다. pip install python-docx")

    try:
        doc = Document(file_path)
        lines = []

        # 문단 추출
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                lines.append(text)

        # 표 추출
        for table_idx, table in enumerate(doc.tables, start=1):
            lines.append(f"\n[Table {table_idx}]")

            for row_idx, row in enumerate(table.rows, start=1):
                cells = []

                for cell in row.cells:
                    value = cell.text.strip().replace("\n", " ")
                    if value:
                        cells.append(value)

                if cells:
                    lines.append(f"Row {row_idx}: " + " | ".join(cells))

        return ExtractedText(
            text="\n".join(lines).strip(),
            source_type="docx",
            parse_method="docx_paragraphs_tables",
        )

    except Exception as e:
        raise HTTPException(500, f"DOCX 텍스트 추출 실패: {e}")


def extract_xlsx(file_path: str) -> ExtractedText:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise HTTPException(500, "openpyxl 라이브러리가 필요합니다. pip install openpyxl")

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"\n[Sheet: {sheet_name}]")

            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                values = []

                for cell in row:
                    if cell is None:
                        continue

                    value = str(cell).strip()
                    if value:
                        values.append(value)

                if values:
                    lines.append(f"Row {row_idx}: " + " | ".join(values))

        wb.close()

        return ExtractedText(
            text="\n".join(lines).strip(),
            source_type="xlsx",
            parse_method="openpyxl_values",
        )

    except Exception as e:
        raise HTTPException(500, f"XLSX 텍스트 추출 실패: {e}")


def extract_hwpx(file_path: str) -> ExtractedText:
    try:
        lines = []

        with zipfile.ZipFile(file_path, "r") as zf:
            xml_names = [
                name for name in zf.namelist()
                if name.lower().startswith("contents/section") and name.lower().endswith(".xml")
            ]

            if not xml_names:
                xml_names = [
                    name for name in zf.namelist()
                    if name.lower().endswith(".xml")
                ]

            for name in sorted(xml_names):
                try:
                    xml_bytes = zf.read(name)
                    root = ET.fromstring(xml_bytes)

                    texts = []
                    for elem in root.iter():
                        # HWPX의 텍스트 노드는 보통 hp:t 형태지만,
                        # namespace가 붙기 때문에 endswith로 처리
                        if elem.tag.endswith("}t") or elem.tag.endswith("t"):
                            if elem.text and elem.text.strip():
                                texts.append(elem.text.strip())

                    if texts:
                        lines.append(f"\n[Section: {name}]")
                        lines.append(" ".join(texts))

                except Exception:
                    continue

        text = "\n".join(lines).strip()

        if not text:
            raise HTTPException(400, "HWPX에서 추출된 텍스트가 없습니다")

        return ExtractedText(
            text=text,
            source_type="hwpx",
            parse_method="zip_xml",
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"HWPX 텍스트 추출 실패: {e}")


def extract_hwp(file_path: str) -> ExtractedText:
    """
    HWP는 문서 버전과 구조에 따라 안정성이 낮음.
    pyhwp 설치 시 hwp5txt 명령어를 이용해 텍스트 추출을 시도.
    """
    try:
        result = subprocess.run(
            ["hwp5txt", file_path],
            capture_output=True,
            text=True,
            timeout=30,
            encoding="utf-8",
            errors="ignore",
        )

        if result.returncode != 0:
            raise HTTPException(
                400,
                "HWP 텍스트 추출 실패. HWPX, PDF, DOCX로 변환 후 업로드해주세요."
            )

        text = result.stdout.strip()

        if not text:
            raise HTTPException(
                400,
                "HWP에서 추출된 텍스트가 없습니다. HWPX, PDF, DOCX로 변환 후 업로드해주세요."
            )

        return ExtractedText(
            text=text,
            source_type="hwp",
            parse_method="hwp5txt",
        )

    except FileNotFoundError:
        raise HTTPException(
            500,
            "hwp5txt 명령어가 없습니다. pyhwp 설치 또는 HWP 변환 도구가 필요합니다."
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"HWP 텍스트 추출 실패: {e}")


def extract_image_with_ocr(file_path: str) -> ExtractedText:
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        raise HTTPException(
            500,
            "OCR 사용을 위해 pillow, pytesseract 라이브러리가 필요합니다"
        )

    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang="kor+eng")

        return ExtractedText(
            text=text or "",
            source_type="image",
            parse_method="ocr",
        )

    except Exception as e:
        raise HTTPException(500, f"이미지 OCR 실패: {e}")


def extract_pdf_with_ocr(file_path: str) -> ExtractedText:
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        raise HTTPException(
            500,
            "PDF OCR 사용을 위해 pdf2image, pytesseract 라이브러리가 필요합니다"
        )

    try:
        images = convert_from_path(file_path, dpi=200)
        texts = []

        for page_no, img in enumerate(images, start=1):
            txt = pytesseract.image_to_string(img, lang="kor+eng")
            if txt and txt.strip():
                texts.append(f"[Page {page_no}]\n{txt.strip()}")

        merged = "\n\n".join(texts).strip()

        return ExtractedText(
            text=merged,
            source_type="pdf",
            parse_method="ocr",
            page_count=len(images),
        )

    except Exception as e:
        raise HTTPException(500, f"PDF OCR 실패: {e}")